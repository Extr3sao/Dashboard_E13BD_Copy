import re
import docx
import os
from typing import Dict, List

class FileParser:
    @staticmethod
    def parse_sql_file(file_path: str) -> Dict[str, str]:
        """
        Parses a TXT file containing multiple SQL queries separated by -- name: <name>
        """
        if not os.path.exists(file_path):
            return {}

        queries = {}
        current_name = None
        current_sql = []

        with open(file_path, 'r', encoding='utf-8') as f:
            for line in f:
                # Check for query name header
                name_match = re.match(r'^\s*--\s*name:\s*(\w+)', line, re.IGNORECASE)
                if name_match:
                    if current_name and current_sql:
                        queries[current_name] = " ".join(current_sql).strip()
                    current_name = name_match.group(1)
                    current_sql = []
                elif current_name:
                    # Filter out comments and empty lines inside the query
                    clean_line = re.sub(r'--.*$', '', line).strip()
                    if clean_line:
                        current_sql.append(clean_line)
            
            # Add last query
            if current_name and current_sql:
                queries[current_name] = " ".join(current_sql).strip()

        return queries

    @staticmethod
    def parse_cleaning_manual(file_path: str) -> Dict:
        """
        Parses the Word document containing cleaning rules.
        For now, it extracts text and simple table data.
        """
        if not os.path.exists(file_path):
            return {"error": "File not found"}

        doc = docx.Document(file_path)
        content = {
            "text": [],
            "tables": []
        }

        for para in doc.paragraphs:
            if para.text.strip():
                content["text"].append(para.text.strip())

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            content["tables"].append(table_data)

        return content

if __name__ == "__main__":
    # Test Parser
    parser = FileParser()
    # Mock SQL test would go here
    print("Parser initialized.")
